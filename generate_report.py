"""
Generate a Word document report for the Ag-Mg Phase Diagram project.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table-cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_line(doc):
    """Insert a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1E3A5F')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)


def styled_heading(doc, text, level=1, color=RGBColor(0x1E, 0x3A, 0x5F)):
    """Add a heading with custom colour."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def body_para(doc, text, bold_prefix=None, space_after=6):
    """Add a body paragraph, optionally with a bold prefix."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        p.add_run(' ')
    p.add_run(text)
    return p


def bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix + ': ')
        r.bold = True
    p.add_run(text)
    return p


# ─── Document Setup ─────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

# Default body font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Heading colours
for i, col in [(1, RGBColor(0x1E, 0x3A, 0x5F)),
               (2, RGBColor(0x1E, 0x5F, 0x99)),
               (3, RGBColor(0x00, 0x70, 0xC0))]:
    try:
        h_style = doc.styles[f'Heading {i}']
        h_style.font.color.rgb = col
        h_style.font.name = 'Calibri'
    except KeyError:
        pass

# ─── Cover / Title ──────────────────────────────────────────────────────────

doc.add_paragraph()   # spacer

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Ag–Mg Interactive Binary Phase Diagram")
title_run.bold = True
title_run.font.size = Pt(26)
title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
title_run.font.name = 'Calibri'

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run(
    "A Browser-Based Interactive Tool for Materials Science Education"
)
sub_run.font.size = Pt(14)
sub_run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
sub_run.font.name = 'Calibri'
sub_p.paragraph_format.space_after = Pt(6)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run(f"April 2026  ·  Packathon Project Presentation")
date_run.font.size = Pt(11)
date_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
date_run.font.name = 'Calibri Italic'

doc.add_paragraph()   # spacer
add_horizontal_line(doc)
doc.add_paragraph()   # spacer

# ─── 1. Project Overview ─────────────────────────────────────────────────────

styled_heading(doc, "1. Project Overview", level=1)

body_para(doc,
    "The Ag–Mg Interactive Binary Phase Diagram is a zero-dependency, "
    "fully browser-based scientific tool that allows users to explore the "
    "Silver–Magnesium (Ag–Mg) binary alloy system in real time. "
    "It was developed as a Packathon project, aiming to make advanced "
    "materials science concepts — phase boundaries, lever-rule calculations, "
    "and microstructure evolution — accessible through an elegant, interactive interface.")

body_para(doc,
    "The application runs entirely inside a single HTML file (≈ 80 KB) "
    "with no external libraries or build steps required. "
    "It can be opened directly in any modern web browser or hosted on a "
    "static platform such as Vercel.")

add_horizontal_line(doc)

# ─── 2. Motivation & Problem Statement ────────────────────────────────────────

styled_heading(doc, "2. Motivation & Problem Statement", level=1)

body_para(doc,
    "Understanding binary phase diagrams is fundamental to metallurgy and "
    "materials engineering, yet existing digital tools are often: heavy "
    "(requiring MATLAB or specialised CALPHAD software), not interactive, "
    "lacking microstructure visualisations, or behind expensive paywalls.")

body_para(doc,
    "Our goal was to create a lightweight, freely accessible, visually rich "
    "educational tool that:")
for item in [
    "Renders the full Ag–Mg phase diagram using literature-backed data.",
    "Allows real-time exploration of phase states at any (T, C₀) point.",
    "Applies the lever rule automatically for two-phase regions.",
    "Visualises representative microstructures for each phase state.",
    "Exports analysis as a PDF — directly from the browser.",
]:
    bullet(doc, item)

add_horizontal_line(doc)

# ─── 3. Key Features ──────────────────────────────────────────────────────────

styled_heading(doc, "3. Key Features", level=1)

features = [
    ("Interactive Phase Diagram",
     "A zoomable, hoverable Canvas-rendered phase diagram covering "
     "100–1,000 °C and 0–100 wt.% Mg with colour-coded phase regions."),
    ("Real-Time Controls",
     "Temperature and composition sliders instantly update the phase state, "
     "lever-rule display, and microstructure canvas."),
    ("Lever Rule Engine",
     "Automatic calculation of phase fractions (f₁, f₂) for any two-phase "
     "field, displayed as a segmented bar chart."),
    ("Microstructure Visualisation",
     "A Canvas-based rendering engine draws representative microstructures: "
     "uniform α grains, precipitates, dendrites, lamellar eutectic structures, "
     "and intermetallic patterns — all generated algorithmically."),
    ("Temperature Sweep Animation",
     "An animated sweep moves the temperature marker across the diagram at "
     "adjustable speed, showing how phases evolve on cooling or heating."),
    ("Phase Fraction Pie Chart",
     "A real-time pie chart shows proportional phase fractions."),
    ("Hover Tooltips",
     "Hovering over any point on the phase diagram shows exact phase "
     "information for that coordinate."),
    ("PDF Export",
     "A single click opens the browser print dialog for a print-ready "
     "analysis report."),
]

for name, desc in features:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(name + ': ')
    r.bold = True
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    p.add_run(desc)

add_horizontal_line(doc)

# ─── 4. Ag–Mg Alloy System ────────────────────────────────────────────────────

styled_heading(doc, "4. The Ag–Mg Alloy System — Scientific Background", level=1)

body_para(doc,
    "Silver (Ag) and Magnesium (Mg) form a rich binary system with multiple "
    "intermetallic compounds and invariant reactions. Key thermodynamic points "
    "used in this project are listed below.")

# System-constants table
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Parameter'
hdr[1].text = 'Value'
for cell in hdr:
    set_cell_bg(cell, '1E3A5F')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

data = [
    ("Ag melting point",            "961.93 °C"),
    ("Mg melting point",            "650 °C"),
    ("Eutectic E2 (γ + δ + L)",     "~472 °C, ~50 wt.% Mg"),
    ("Peritectic point P",          "~492 °C, ~32 wt.% Mg"),
    ("β' congruent maximum",        "~820 °C, ~20 wt.% Mg"),
    ("Metatectic E1 (α + β' + L)",  "~760 °C, ~10 wt.% Mg"),
    ("Liquidus data source",        "Experimental CSV (9,986 → 276 points after smoothing)"),
    ("Accuracy (piecewise model)",  "±10 °C  ·  ±2 at.%"),
]

for row_data in data:
    row = table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]

doc.add_paragraph()  # spacer after table

body_para(doc,
    "Phases present in the Ag–Mg system: liquid (L), α-Ag solid solution, "
    "β' (B2-ordered AgMg intermetallic), γ (Ag₃Mg), δ (AgMg₃ / ε), "
    "and pure Mg solid solution. The intermetallics β' and γ occupy "
    "narrow composition windows requiring careful piecewise boundary definitions.")

add_horizontal_line(doc)

# ─── 5. How It Was Built ──────────────────────────────────────────────────────

styled_heading(doc, "5. How the Tool Was Built", level=1)

# 5.1 Technology Stack
styled_heading(doc, "5.1 Technology Stack", level=2)

stack = [
    ("HTML5 + CSS3",
     "All UI structure and styling — dark-mode design system with CSS custom "
     "properties, glassmorphism cards, animated gradients, and responsive grid layout."),
    ("Vanilla JavaScript (ES2020)",
     "All application logic, data processing, phase detection, and Canvas "
     "rendering — zero runtime dependencies."),
    ("HTML5 Canvas 2D API",
     "Two canvas elements: one for the phase diagram itself, and one for the "
     "microstructure visualisation panel."),
    ("Google Fonts (Inter + JetBrains Mono)",
     "Professional typography loaded at runtime for readability."),
    ("Python (pre-processing only)",
     "process_liquidus.py and process_alpha_solidus.py were used offline to "
     "parse and smooth raw experimental CSV datasets into compact JavaScript arrays."),
    ("Vercel",
     "Static-site deployment — the single index.html is served globally via CDN."),
]

for name, desc in stack:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(name + ': ')
    r.bold = True
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    p.add_run(desc)

# 5.2 Phase Data
styled_heading(doc, "5.2 Phase Boundary Data Acquisition", level=2)

body_para(doc,
    "Phase boundary data was sourced from the ASM Binary Alloy Phase Diagrams "
    "(Massalski, 1990) and ASM Handbook Vol. 3 (1992). "
    "Raw experimental liquidus data was provided as a CSV file containing "
    "9,986 data points. Before embedding this data into the application, "
    "a Python pre-processing pipeline was written:")

for step in [
    "Load the CSV with pandas.",
    "Sort and remove duplicate (composition, temperature) pairs.",
    "Apply a Savitzky–Golay smoothing filter to reduce noise.",
    "Downsample to ~276 representative points using uniform step sampling.",
    "Output the result as a compact JavaScript array literal.",
]:
    bullet(doc, step)

body_para(doc,
    "The α-solidus boundary was processed similarly from alpha solidus.csv. "
    "All other boundaries (solvus lines, intermetallic regions) were digitised "
    "by hand as piecewise-linear [wt.% Mg, T °C] arrays directly in JavaScript.")

# 5.3 Phase Detection
styled_heading(doc, "5.3 Phase Detection Algorithm", level=2)

body_para(doc,
    "Given a user-selected (temperature T, composition C₀), the application "
    "determines the current phase state using a boundary-based algorithm:")

for step in [
    "If T > liquidus(C₀) → Liquid.",
    "Test each predefined two-phase region by checking if C₀ falls between "
    "the two boundary curves at temperature T.",
    "Test single-phase solid regions (α, β', γ, δ, Mg-ss) using solvus curves.",
    "Check invariant temperatures (eutectic 472 °C, peritectic 492 °C) within ±5 °C tolerance.",
    "Return phase label, description, and the boundary concentrations C₁ / C₂ "
    "needed for the lever rule.",
]:
    bullet(doc, step)

# 5.4 Lever Rule
styled_heading(doc, "5.4 Lever Rule Calculation", level=2)

body_para(doc,
    "For a composition C₀ in a two-phase field [C₁, C₂], the application "
    "calculates the mass fractions using the classical lever rule:")

eq_p = doc.add_paragraph()
eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq_run = eq_p.add_run(
    "f₁ = (C₂ − C₀) / (C₂ − C₁)          f₂ = (C₀ − C₁) / (C₂ − C₁)"
)
eq_run.font.name = 'Courier New'
eq_run.font.size = Pt(11)
eq_run.font.bold = True
eq_run.font.color.rgb = RGBColor(0x1E, 0x5F, 0x99)
eq_p.paragraph_format.space_before = Pt(4)
eq_p.paragraph_format.space_after = Pt(8)

body_para(doc,
    "Results are displayed as: numerical fractions in the Lever Rule panel, "
    "an animated segmented bar chart, and updated microstructure proportions.")

# 5.5 Microstructure Visualisation
styled_heading(doc, "5.5 Microstructure Visualisation Engine", level=2)

body_para(doc,
    "A key innovation of this project is the algorithmically generated "
    "microstructure visualisation. Each phase state is mapped to a specific "
    "rendering mode drawn on an HTML Canvas:")

micro_table = doc.add_table(rows=1, cols=3)
micro_table.style = 'Table Grid'
hdr = micro_table.rows[0].cells
for cell, text in zip(hdr, ['Phase State', 'Microstructure Type', 'Visual Pattern']):
    cell.text = text
    set_cell_bg(cell, '1E3A5F')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

micro_data = [
    ("Liquid",               "Liquid",        "Animated swirling circles / bubbles"),
    ("α solid solution",     "Uniform grains","Voronoi-like polygonal grain structure"),
    ("β' intermetallic",     "Intermetallic", "Regular ordered rectangular tiles"),
    ("α + β' two-phase",     "Precipitate",   "Grain matrix with second-phase particles"),
    ("Eutectic (L↔α+β')",   "Lamellar",      "Alternating dark/light lamellae bands"),
    ("Peritectic",           "Dendritic",     "Branched dendritic solidification arms"),
    ("γ, δ regions",         "Intermetallic", "Dense ordered tiling pattern"),
]

for row_data in micro_data:
    row = micro_table.add_row().cells
    for cell, text in zip(row, row_data):
        cell.text = text

doc.add_paragraph()

body_para(doc,
    "All microstructure patterns use a seeded pseudo-random number generator "
    "so results are deterministic and reproducible for any given (T, C₀). "
    "The rendering is triggered every time the user moves a slider or clicks "
    "the diagram, providing instantaneous visual feedback.")

# 5.6 UI Design
styled_heading(doc, "5.6 UI & Design System", level=2)

body_para(doc,
    "The interface was designed with a dark-mode aesthetic inspired by "
    "modern developer tools and scientific dashboards. "
    "The design system is defined entirely in CSS custom properties and includes:")

for item in [
    "A deep-dark background (#050505) with layered surface cards.",
    "Purple-to-teal gradient accents (--accent: #8b5cf6, --teal: #14b8a6).",
    "Animated conic-gradient glow on the Microstructure card.",
    "Smooth cubic-bezier transitions on all interactive elements.",
    "A fully responsive CSS Grid layout (sidebar + main canvas) that stacks on mobile.",
    "Custom-styled range sliders with gradient thumbs.",
    "Hover-activated tooltips positioned via JavaScript mouse tracking.",
]:
    bullet(doc, item)

add_horizontal_line(doc)

# ─── 6. Development Timeline ──────────────────────────────────────────────────

styled_heading(doc, "6. Development Journey", level=1)

styled_heading(doc, "Phase 1 — Initial Prototype", level=2)
body_para(doc,
    "The project began with a basic phase diagram canvas rendering experiment "
    "using approximate boundary data. Temperature and composition sliders "
    "were added and a simple phase-label display was implemented.")

styled_heading(doc, "Phase 2 — Data-Driven Accuracy Upgrade", level=2)
body_para(doc,
    "The approximate boundary curves were replaced with literature-backed "
    "experimental data from the ASM Handbook. Python scripts (process_liquidus.py, "
    "process_alpha_solidus.py) were written to pre-process 9,986-point CSV datasets "
    "into compact, browser-embeddable JavaScript arrays.")

styled_heading(doc, "Phase 3 — Phase Detection Overhaul", level=2)
body_para(doc,
    "The phase detection engine was rewritten using a boundary-traversal approach "
    "to correctly identify narrow regions (β', γ) that required strict "
    "piecewise boundary enforcement. Edge cases at invariant temperatures "
    "(eutectic 472 °C, peritectic 492 °C) were handled with tolerance bands.")

styled_heading(doc, "Phase 4 — Lever Rule & Visualisation", level=2)
body_para(doc,
    "The lever rule calculation module was integrated, outputting phase fractions "
    "to both a segmented bar chart and a displayed data panel. "
    "The composition sweep animation and PDF export feature were added.")

styled_heading(doc, "Phase 5 — Microstructure Visualisation", level=2)
body_para(doc,
    "A complete Canvas-based microstructure rendering engine was designed and "
    "implemented. Each phase state was mapped to a distinct visual pattern "
    "(uniform grains, precipitates, lamellar, dendritic, intermetallic, liquid). "
    "The Microstructure card was given a glowing animated border to make it "
    "visually distinctive.")

styled_heading(doc, "Phase 6 — Polish & Deployment", level=2)
body_para(doc,
    "Final UI polish pass: responsive layout fixes, tooltip improvements, "
    "system-constants table, and README / documentation. "
    "The application was deployed to Vercel as a zero-config static site.")

add_horizontal_line(doc)

# ─── 7. Assumptions & Limitations ────────────────────────────────────────────

styled_heading(doc, "7. Assumptions & Limitations", level=1)

for item in [
    "Phase boundaries are represented as piecewise-linear curves, not full CALPHAD thermodynamic polynomials.",
    "The model assumes equilibrium conditions only; kinetics (e.g., undercooling, rapid solidification) are not modelled.",
    "Accuracy is approximately ±10 °C and ±2 at.% compared to reference phase diagrams.",
    "Metastable phases are excluded from the analysis.",
    "Microstructure images are schematic and algorithmically generated — they are representative, not simulation-accurate.",
    "Composition is expressed in wt.% Mg (weight percent) throughout the tool.",
]:
    bullet(doc, item)

add_horizontal_line(doc)

# ─── 8. File Structure ────────────────────────────────────────────────────────

styled_heading(doc, "8. Project File Structure", level=1)

files = [
    ("index.html",               "Complete single-file web application (~2,100 lines)"),
    ("Final Website.html",       "Backup / alternate version of the final build"),
    ("Liquidus Dataset.csv",     "Raw experimental liquidus data (9,986 data points)"),
    ("alpha solidus.csv",        "Raw α-solidus boundary data"),
    ("process_liquidus.py",      "Python pre-processing script for liquidus CSV"),
    ("process_alpha_solidus.py", "Python pre-processing script for α-solidus CSV"),
    ("liquidus_js.txt",          "Processed JS array output of liquidus data"),
    ("alpha_solidus_js.txt",     "Processed JS array output of α-solidus data"),
    ("isotherm.py",              "Utility script for isotherm analysis"),
    ("Ag Mg phase diagram.png",  "Reference phase diagram image"),
    ("REPORT.md",                "Technical methodology report (Markdown)"),
    ("README.md",                "Project overview and usage guide"),
    ("package.json",             "Project metadata for npm/Vercel"),
    ("vercel.json",              "Vercel deployment configuration"),
]

file_table = doc.add_table(rows=1, cols=2)
file_table.style = 'Table Grid'
hdr = file_table.rows[0].cells
for cell, text in zip(hdr, ['File', 'Purpose']):
    cell.text = text
    set_cell_bg(cell, '1E3A5F')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for row_data in files:
    row = file_table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    # make filename monospace
    for run in row[0].paragraphs[0].runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9.5)

doc.add_paragraph()

add_horizontal_line(doc)

# ─── 9. References ────────────────────────────────────────────────────────────

styled_heading(doc, "9. References", level=1)

refs = [
    "Massalski, T. B. (Ed.). (1990). Binary Alloy Phase Diagrams (2nd ed.). ASM International.",
    "ASM International. (1992). ASM Handbook Volume 3: Alloy Phase Diagrams. ASM International.",
    "CALPHAD: Calculation of Phase Diagrams — A Comprehensive Guide. (General methodology reference.)",
    "HTML5 Canvas 2D API — MDN Web Docs. https://developer.mozilla.org/en-US/docs/Web/API/Canvas_API",
    "Inter Typeface — Rasmus Andersson. https://rsms.me/inter/",
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(f"[{i}] ").bold = True
    p.add_run(ref)

add_horizontal_line(doc)

# ─── Footer note ─────────────────────────────────────────────────────────────

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_p.add_run(
    "Ag–Mg Phase Diagram Tool  ·  Packathon 2026  ·  "
    "Built with HTML5 Canvas + Vanilla JavaScript  ·  Zero Dependencies"
)
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
footer_run.font.italic = True

# ─── Save ─────────────────────────────────────────────────────────────────────

output_path = r"c:\Users\RUDRA SHARMA\OneDrive\Desktop\Packathon\ag-mg-phase-diagram\Ag-Mg Phase Diagram Project Report.docx"
doc.save(output_path)
print(f"Document saved -> {output_path}")
